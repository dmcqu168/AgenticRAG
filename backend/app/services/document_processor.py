import os
import io
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any
import logging
from datetime import datetime

# Document processing libraries
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from docx import Document as DocxDocument
    import docx2txt
    import pandas as pd
    import pytesseract
    from PIL import Image
    import numpy as np
    import cv2
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LAParams, LTTextContainer, LTFigure
    import pdf2image
    
    # Configure Tesseract path (update this if Tesseract is installed in a different location)
    pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'  # Default Homebrew location
    
    # Set TESSDATA_PREFIX environment variable if needed
    os.environ['TESSDATA_PREFIX'] = '/opt/homebrew/share/tessdata/'
    
except ImportError as e:
    logger = logging.getLogger(__name__)
    logger.warning(f"Missing required OCR dependencies: {str(e)}")
    pass

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of different document types for the RAG system"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize the document processor
        
        Args:
            chunk_size: Number of characters per chunk
            chunk_overlap: Number of overlapping characters between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_text,
            '.csv': self._process_csv,
        }
    
    async def process(self, file_path: Union[str, Path], metadata: Optional[Dict] = None) -> List[Dict]:
        """Process a document and return chunks with metadata
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata to include with each chunk
            
        Returns:
            List of document chunks with metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract text from the document
        try:
            text = await self.supported_formats[file_extension](file_path)
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            raise ValueError(f"Failed to process document: {str(e)}")
        
        # Clean and chunk the text
        chunks = self._chunk_text(text)
        
        # Prepare metadata for each chunk
        base_metadata = {
            'source': str(file_path.name),
            'file_type': file_extension[1:],  # Remove the dot
            'processing_date': datetime.utcnow().isoformat(),
            **(metadata or {})
        }
        
        # Add chunk-specific metadata
        result = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **base_metadata,
                'chunk_id': f"{file_path.stem}_{i}",
                'chunk_index': i,
                'total_chunks': len(chunks)
            }
            result.append({
                'content': chunk,
                'metadata': chunk_metadata
            })
            
        return result
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
            
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunks.append(text[start:end])
            
            if end == text_length:
                break
                
            # Move back by overlap amount, but not past the start of the last chunk
            start = end - min(self.chunk_overlap, self.chunk_size // 2)
            
        return chunks
    
    def _extract_text_with_ocr(self, image: Image.Image) -> str:
        """Extract text from an image using OCR"""
        try:
            # Convert to grayscale for better OCR accuracy
            if image.mode != 'L':
                image = image.convert('L')
            
            # Convert to numpy array for OpenCV processing
            img_array = np.array(image)
            
            # Apply image preprocessing
            # 1. Convert to grayscale if not already
            if len(img_array.shape) > 2:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # 2. Apply thresholding to get binary image
            _, img_array = cv2.threshold(img_array, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # 3. Apply dilation and erosion to remove noise
            kernel = np.ones((1, 1), np.uint8)
            img_array = cv2.dilate(img_array, kernel, iterations=1)
            img_array = cv2.erode(img_array, kernel, iterations=1)
            
            # Convert back to PIL Image
            processed_img = Image.fromarray(img_array)
            
            # Perform OCR with Tesseract
            text = pytesseract.image_to_string(processed_img, lang='eng')
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error in OCR processing: {str(e)}")
            return ""
    
    def _process_pdf_page(self, page) -> str:
        """Process a single PDF page, handling both text and images"""
        text = ""
        
        # First try to extract text directly
        page_text = page.extract_text()
        if page_text and page_text.strip():
            text += page_text + "\n\n"
        
        # Check for images in the page
        if '/XObject' in page['/Resources']:
            xObject = page['/Resources']['/XObject'].get_object()
            
            for obj in xObject:
                if xObject[obj]['/Subtype'] == '/Image':
                    try:
                        # Extract the image data
                        img_data = xObject[obj].get_data()
                        
                        # Try to identify the image format
                        img_format = None
                        if '/Filter' in xObject[obj]:
                            if '/FlateDecode' in xObject[obj]['/Filter']:
                                img_format = 'png'
                            elif '/DCTDecode' in xObject[obj]['/Filter']:
                                img_format = 'jpeg'
                            elif '/JPXDecode' in xObject[obj]['/Filter']:
                                img_format = 'jpeg2000'
                        
                        # Try to open the image with the detected format
                        try:
                            img = Image.open(io.BytesIO(img_data))
                            # Process the image with OCR
                            ocr_text = self._extract_text_with_ocr(img)
                            if ocr_text:
                                text += f"[IMAGE CONTENT - {img_format or 'unknown'} format]\n{ocr_text}\n\n"
                        except Exception as img_open_error:
                            # If direct opening fails, try with PIL's ImageFile
                            from PIL import ImageFile
                            ImageFile.LOAD_TRUNCATED_IMAGES = True
                            try:
                                img = Image.open(io.BytesIO(img_data))
                                # Convert to RGB mode if needed
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    img = img.convert('RGB')
                                # Save as JPEG and reload to ensure compatibility
                                temp_img = io.BytesIO()
                                img.save(temp_img, format='JPEG')
                                img = Image.open(temp_img)
                                # Process the image with OCR
                                ocr_text = self._extract_text_with_ocr(img)
                                if ocr_text:
                                    text += f"[IMAGE CONTENT - converted to JPEG]\n{ocr_text}\n\n"
                            except Exception as convert_error:
                                logger.warning(f"Could not convert image in PDF: {str(convert_error)}")
                                # Try to get basic image info even if we can't process it
                                try:
                                    img_size = f"{xObject[obj].get('Width', '?')}x{xObject[obj].get('Height', '?')}"
                                    text += f"[UNPROCESSED IMAGE - Format: {img_format or 'unknown'}, Size: {img_size}]\n\n"
                                except:
                                    text += "[UNPROCESSED IMAGE - Could not extract details]\n\n"
                    except Exception as img_error:
                        logger.warning(f"Error processing image in PDF: {str(img_error)}")
                        text += "[ERROR PROCESSING IMAGE]\n\n"
        
        return text
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file with OCR for images and scanned pages"""
        try:
            text = ""
            
            # First, try to extract text using PyMuPDF (fitz) which handles more PDF formats
            try:
                import fitz  # PyMuPDF
                
                doc = fitz.open(file_path)
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Extract text from the page
                    page_text = page.get_text()
                    if page_text.strip():
                        text += f"[PAGE {page_num + 1}]\n{page_text}\n\n"
                    
                    # Extract and process images
                    image_list = page.get_images(full=True)
                    if image_list:
                        for img_index, img in enumerate(image_list, 1):
                            try:
                                # Get the image
                                xref = img[0]
                                base_image = doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                
                                # Process the image with OCR
                                with Image.open(io.BytesIO(image_bytes)) as img_pil:
                                    ocr_text = self._extract_text_with_ocr(img_pil)
                                    if ocr_text:
                                        text += f"[IMAGE {img_index} ON PAGE {page_num + 1}]\n{ocr_text}\n\n"
                            except Exception as img_error:
                                logger.warning(f"Error processing image {img_index} on page {page_num + 1}: {str(img_error)}")
                                text += f"[UNPROCESSED IMAGE {img_index} ON PAGE {page_num + 1}]\n\n"
                
                # If we got text, return it
                if text.strip():
                    return text
                    
            except Exception as fitz_error:
                logger.warning(f"PyMuPDF processing failed, falling back to PyPDF2: {str(fitz_error)}")
                text = ""
            
            # Fallback to PyPDF2 if PyMuPDF fails or returns no text
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    # Process each page
                    for page_num, page in enumerate(reader.pages, 1):
                        page_text = self._process_pdf_page(page)
                        if page_text:
                            text += f"[PAGE {page_num} - PDF TEXT]\n{page_text}\n\n"
            except Exception as pypdf_error:
                logger.warning(f"PyPDF2 processing failed: {str(pypdf_error)}")
            
            # If still no text, try OCR on the entire document
            if not text.strip():
                logger.info("No text found with PyMuPDF or PyPDF2, attempting full document OCR...")
                
                # Convert PDF to images
                try:
                    images = convert_from_path(
                        str(file_path),
                        dpi=300,  # Higher DPI for better OCR accuracy
                        thread_count=4,
                        grayscale=True
                    )
                    
                    # Process each page with OCR
                    for i, image in enumerate(images, 1):
                        logger.info(f"Processing page {i}/{len(images)} with OCR...")
                        ocr_text = self._extract_text_with_ocr(image)
                        if ocr_text:
                            text += f"[PAGE {i} - OCR EXTRACT]\n{ocr_text}\n\n"
                except Exception as ocr_error:
                    logger.error(f"Error during OCR processing: {str(ocr_error)}")
            
            if not text.strip():
                logger.warning(f"No text could be extracted from {file_path}")
                return f"[UNABLE TO EXTRACT TEXT FROM DOCUMENT: {file_path.name}]"
                
            return text
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}", exc_info=True)
            return f"[ERROR PROCESSING DOCUMENT: {str(e)}]"
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file, including OCR for images"""
        try:
            # First try to extract text and images using docx2txt
            temp_dir = tempfile.mkdtemp()
            text = docx2txt.process(file_path, temp_dir)
            
            # Process any extracted images with OCR
            if os.path.exists(temp_dir):
                for img_file in os.listdir(temp_dir):
                    if img_file.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
                        try:
                            img_path = os.path.join(temp_dir, img_file)
                            with Image.open(img_path) as img:
                                ocr_text = self._extract_text_with_ocr(img)
                                if ocr_text:
                                    text += f"\n\n[IMAGE CONTENT FROM {img_file}]:\n{ocr_text}"
                        except Exception as img_error:
                            logger.warning(f"Error processing image {img_file}: {str(img_error)}")
                            continue
                
                # Clean up temporary directory
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
            
            # If no text was extracted, try the python-docx method as fallback
            if not text.strip():
                logger.info("No text found with docx2txt, trying python-docx...")
                try:
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text])
                except Exception as e2:
                    logger.warning(f"python-docx extraction failed: {str(e2)}")
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            # Final fallback to python-docx
            try:
                doc = DocxDocument(file_path)
                return "\n".join([para.text for para in doc.paragraphs if para.text])
            except Exception as e2:
                logger.error(f"All DOCX extraction methods failed: {str(e2)}")
                raise e
    
    async def _process_text(self, file_path: Path) -> str:
        """Read text from a plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            raise
    
    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV file, handling tables and any referenced images with OCR"""
        try:
            # Read the CSV file
            df = pd.read_csv(file_path)
            
            # Convert the dataframe to a readable string
            result = []
            
            # Add a header
            result.append(f"=== CSV Content: {file_path.name} ===")
            
            # Add column names
            result.append("\n".join([", ".join(df.columns)]))
            
            # Add a separator
            result.append("-" * 50)
            
            # Add data rows
            for _, row in df.iterrows():
                # Check each cell for potential image references
                row_data = []
                for col in df.columns:
                    cell_value = str(row[col])
                    
                    # Check if the cell might reference an image file
                    if cell_value.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
                        img_path = Path(cell_value)
                        if img_path.is_file():
                            try:
                                with Image.open(img_path) as img:
                                    ocr_text = self._extract_text_with_ocr(img)
                                    if ocr_text:
                                        cell_value = f"[IMAGE: {img_path.name}] {ocr_text}"
                            except Exception as img_error:
                                logger.warning(f"Could not process image {img_path}: {str(img_error)}")
                    
                    row_data.append(cell_value)
                
                result.append(", ".join(row_data))
            
            # Add a summary of the data
            result.append("\n=== Data Summary ===")
            result.append(f"Total Rows: {len(df)}")
            result.append(f"Columns: {', '.join(df.columns)}")
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                result.append("\nNumeric Column Statistics:")
                for col in numeric_cols:
                    try:
                        stats = df[col].describe()
                        result.append(f"{col}: mean={stats['mean']:.2f}, min={stats['min']}, max={stats['max']}")
                    except:
                        continue
            
            return "\n".join(result)
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            # Fallback to simple CSV reading if pandas processing fails
            try:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    return f"[RAW CSV CONTENT]\n{f.read()}"
            except Exception as e2:
                logger.error(f"Failed to read CSV as text: {str(e2)}")
                raise e
